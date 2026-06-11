#!/usr/bin/env python3
"""
DOCX 文本/表格提取脚本
提取 DOCX 内容为 JSON 格式
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装 python-docx")
    print("运行: pip install python-docx")
    sys.exit(1)


def extract_text(doc: Document) -> list:
    """提取所有段落文本"""
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(
                {
                    "text": para.text,
                    "style": para.style.name if para.style else None,
                }
            )
    return paragraphs


def extract_tables(doc: Document) -> list:
    """提取所有表格"""
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(
            {
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data,
            }
        )
    return tables


def extract_docx(docx_path: str) -> dict:
    """提取 DOCX 内容"""
    doc = Document(docx_path)

    result = {
        "file": str(docx_path),
        "paragraphs": extract_text(doc),
        "tables": extract_tables(doc),
        "metadata": {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    }

    # 文档属性
    if doc.core_properties.title:
        result["metadata"]["title"] = doc.core_properties.title
    if doc.core_properties.author:
        result["metadata"]["author"] = doc.core_properties.author

    return result


def main():
    if len(sys.argv) < 2:
        print("用法: python docx_extract.py <document.docx> [output.json]")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not Path(docx_path).exists():
        print(f"错误: 文件不存在: {docx_path}")
        sys.exit(1)

    result = extract_docx(docx_path)

    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"已保存到: {output_path}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
